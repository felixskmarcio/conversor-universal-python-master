#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Conversor Universal de Documentos - Refactored Application
A professional document converter with improved architecture, security, and performance.

Features:
- Modular design with separation of concerns
- Enhanced security validation
- Performance optimizations with caching
- Comprehensive error handling and logging
- Rate limiting and monitoring
"""

import os
import tempfile
from pathlib import Path
from flask import Flask
from flask_cors import CORS

from .config import get_config
from .core.logging_config import setup_logging
from .api.routes import api_bp
from .core.rate_limiter import IPRateLimiter

# Bibliotecas para conversão de documentos
try:
    import pypandoc
except ImportError:
    pypandoc = None

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    Document = None

try:
    import pdfplumber
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
except ImportError:
    pdfplumber = None
    canvas = None

try:
    import markdown
    from markdown.extensions import codehilite, tables, toc
except ImportError:
    markdown = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

class ConversorUniversalMelhorado:
    """Classe principal para conversão de documentos com preservação de estrutura"""
    
    def __init__(self):
        self.formatos_suportados = {
            'pdf': ['.pdf'],
            'docx': ['.docx', '.doc'],
            'txt': ['.txt'],
            'html': ['.html', '.htm'],
            'md': ['.md', '.markdown']
        }
        
        self.temp_dir = tempfile.mkdtemp()
        
        # Padrões para detecção de estrutura acadêmica
        self.padroes_instituicao = [
            r'^UNIVERSIDADE\s+.*$',
            r'^CENTRO\s+.*$',
            r'^INSTITUTO\s+.*$',
            r'^FACULDADE\s+.*$',
            r'^ESCOLA\s+.*$'
        ]
        
        self.secoes_especiais = {
            'RESUMO', 'ABSTRACT', 'SUMÁRIO', 'ÍNDICE', 'AGRADECIMENTOS',
            'DEDICATÓRIA', 'EPÍGRAFE', 'LISTA DE FIGURAS', 'LISTA DE TABELAS',
            'LISTA DE ABREVIATURAS', 'LISTA DE SIGLAS', 'REFERÊNCIAS',
            'BIBLIOGRAFIA', 'ANEXOS', 'APÊNDICES', 'CONCLUSÃO', 'CONSIDERAÇÕES FINAIS'
        }

    def __del__(self):
        """Limpa arquivos temporários"""
        if hasattr(self, 'temp_dir') and os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def detectar_formato(self, arquivo_path: str) -> str:
        """Detecta o formato do arquivo baseado na extensão e conteúdo"""
        extensao = Path(arquivo_path).suffix.lower()
        
        # Primeiro tenta pela extensão
        for formato, extensoes in self.formatos_suportados.items():
            if extensao in extensoes:
                return formato
        
        # Se não encontrou pela extensão, tenta detectar pelo conteúdo
        try:
            with open(arquivo_path, 'rb') as f:
                primeiros_bytes = f.read(1024)
                
            # Detecta PDF pelo cabeçalho
            if primeiros_bytes.startswith(b'%PDF'):
                return 'pdf'
            
            # Detecta DOCX (arquivo ZIP com estrutura específica)
            if primeiros_bytes.startswith(b'PK\x03\x04'):
                return 'docx'
            
            # Tenta detectar HTML
            try:
                texto = primeiros_bytes.decode('utf-8', errors='ignore').lower()
                if '<html' in texto or '<!doctype html' in texto:
                    return 'html'
            except:
                pass
                
        except Exception:
            pass
        
        raise ValueError(f"Formato não suportado ou não reconhecido: {extensao}")

    def ler_pdf(self, arquivo_path: str) -> str:
        """Extrai texto de arquivo PDF com melhor formatação"""
        if not pdfplumber:
            raise ImportError("pdfplumber não está instalado")
        
        texto = ""
        with pdfplumber.open(arquivo_path) as pdf:
            for pagina in pdf.pages:
                texto_pagina = pagina.extract_text()
                if texto_pagina:
                    texto += texto_pagina + "\n"
        
        return self._validar_e_limpar_texto(texto)

    def ler_docx(self, arquivo_path: str) -> str:
        """Extrai texto de arquivo DOCX preservando estrutura"""
        if not Document:
            raise ImportError("python-docx não está instalado")
        
        doc = Document(arquivo_path)
        texto = ""
        
        # Extrai texto preservando a estrutura de parágrafos
        for paragrafo in doc.paragraphs:
            if paragrafo.text.strip():
                # Preserva formatação de títulos baseada no estilo
                if paragrafo.style.name.startswith('Heading'):
                    texto += f"\n{paragrafo.text}\n"
                else:
                    texto += paragrafo.text + "\n"
        
        return self._validar_e_limpar_texto(texto.strip())

    def ler_txt(self, arquivo_path: str) -> str:
        """Lê arquivo de texto simples"""
        try:
            with open(arquivo_path, 'r', encoding='utf-8') as arquivo:
                texto = arquivo.read()
        except UnicodeDecodeError:
            # Fallback para outras codificações
            with open(arquivo_path, 'r', encoding='latin-1') as arquivo:
                texto = arquivo.read()
        
        return self._validar_e_limpar_texto(texto)

    def ler_html(self, arquivo_path: str) -> str:
        """Extrai texto de arquivo HTML preservando estrutura"""
        if not BeautifulSoup:
            raise ImportError("beautifulsoup4 não está instalado")
        
        with open(arquivo_path, 'r', encoding='utf-8') as arquivo:
            soup = BeautifulSoup(arquivo.read(), 'html.parser')
        
        # Remove scripts e estilos
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Extrai texto preservando quebras de linha em elementos estruturais
        texto = ""
        for elemento in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'div']):
            if elemento.get_text(strip=True):
                texto += elemento.get_text() + "\n"
        
        return self._validar_e_limpar_texto(texto)

    def ler_md(self, arquivo_path: str) -> str:
        """Lê arquivo Markdown"""
        with open(arquivo_path, 'r', encoding='utf-8') as arquivo:
            return arquivo.read()
    
    def _corrigir_espacamento(self, texto: str) -> str:
        """Corrige problemas de espaçamento e formatação no texto extraído"""
        import re
        
        # Remove múltiplas quebras de linha consecutivas
        texto = re.sub(r'\n{3,}', '\n\n', texto)
        
        # Corrige palavras concatenadas (detecta mudança de minúscula para maiúscula)
        texto = re.sub(r'([a-zàáâãäåæçèéêëìíîïñòóôõöøùúûüý])([A-ZÀÁÂÃÄÅÆÇÈÉÊËÌÍÎÏÑÒÓÔÕÖØÙÚÛÜÝ])', r'\1 \2', texto)
        
        # Corrige espaçamento após pontuação
        texto = re.sub(r'([.!?])([A-ZÀÁÂÃÄÅÆÇÈÉÊËÌÍÎÏÑÒÓÔÕÖØÙÚÛÜÝ])', r'\1 \2', texto)
        
        # Corrige espaçamento em números e texto
        texto = re.sub(r'(\d)([A-Za-zÀ-ÿ])', r'\1 \2', texto)
        texto = re.sub(r'([A-Za-zÀ-ÿ])(\d)', r'\1 \2', texto)
        
        # Remove espaços extras
        texto = re.sub(r' {2,}', ' ', texto)
        
        # Corrige quebras de linha em meio a palavras
        texto = re.sub(r'([a-zàáâãäåæçèéêëìíîïñòóôõöøùúûüý])-\n([a-zàáâãäåæçèéêëìíîïñòóôõöøùúûüý])', r'\1\2', texto)
        
        return texto.strip()
    
    def _validar_e_limpar_texto(self, texto: str) -> str:
        """Valida e limpa o texto convertido para melhorar a qualidade"""
        if not texto or not texto.strip():
            return "Documento vazio ou não foi possível extrair texto."
        
        # Aplica correção de espaçamento
        texto = self._corrigir_espacamento(texto)
        
        # Remove caracteres de controle inválidos
        texto = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', texto)
        
        # Corrige encoding de caracteres especiais comuns
        replacements = {
            'â€™': "'",
            'â€œ': '"',
            'â€\x9d': '"',
            'â€\x93': '–',
            'â€\x94': '—',
            'Ã¡': 'á',
            'Ã©': 'é',
            'Ã­': 'í',
            'Ã³': 'ó',
            'Ãº': 'ú',
            'Ã§': 'ç',
            'Ã ': 'à'
        }
        
        for old, new in replacements.items():
            texto = texto.replace(old, new)
        
        return texto
    
    def _detectar_estrutura_documento(self, texto: str) -> list:
        """Detecta a estrutura do documento (títulos, listas, parágrafos, etc.)"""
        linhas = texto.split('\n')
        estrutura = []
        
        for i, linha in enumerate(linhas):
            linha_limpa = linha.strip()
            if not linha_limpa:
                continue
            
            # Detecta instituição (primeira linha em maiúsculas)
            if i < 5 and any(re.match(padrao, linha_limpa, re.IGNORECASE) for padrao in self.padroes_instituicao):
                estrutura.append({'tipo': 'instituicao', 'texto': linha_limpa})
                continue
            
            # Detecta título principal (linha centralizada ou em maiúsculas no início)
            if i < 10 and (linha_limpa.isupper() or len(linha_limpa) > 10) and not any(char.isdigit() for char in linha_limpa[:5]):
                if linha_limpa.upper() in self.secoes_especiais:
                    estrutura.append({'tipo': 'secao_especial', 'texto': linha_limpa})
                elif self._eh_titulo_principal(linha_limpa):
                    estrutura.append({'tipo': 'titulo_principal', 'texto': linha_limpa})
                else:
                    estrutura.append({'tipo': 'titulo', 'texto': linha_limpa})
                continue
            
            # Detecta seções especiais
            if linha_limpa.upper() in self.secoes_especiais:
                estrutura.append({'tipo': 'secao_especial', 'texto': linha_limpa})
                continue
            
            # Detecta títulos e subtítulos
            if self._eh_titulo(linha_limpa):
                estrutura.append({'tipo': 'titulo', 'texto': linha_limpa})
            elif self._eh_subtitulo(linha_limpa):
                estrutura.append({'tipo': 'subtitulo', 'texto': linha_limpa})
            # Detecta listas numeradas
            elif re.match(r'^\d+[.)\s]', linha_limpa):
                estrutura.append({'tipo': 'lista_numerada', 'texto': linha_limpa})
            # Detecta listas com marcadores
            elif re.match(r'^[•\-\*]\s', linha_limpa):
                estrutura.append({'tipo': 'lista_marcador', 'texto': linha_limpa})
            # Detecta citações
            elif linha_limpa.startswith('"') or linha_limpa.startswith('"'):
                estrutura.append({'tipo': 'citacao', 'texto': linha_limpa})
            # Detecta referências bibliográficas
            elif re.match(r'^[A-Z][A-Z\s,]+\d{4}', linha_limpa):
                estrutura.append({'tipo': 'referencia', 'texto': linha_limpa})
            else:
                estrutura.append({'tipo': 'paragrafo', 'texto': linha_limpa})
        
        return estrutura
    
    def _eh_titulo_principal(self, linha: str) -> bool:
        """Verifica se a linha é um título principal"""
        linha = linha.strip()
        return (
            len(linha) > 5 and
            len(linha) < 100 and
            not linha.endswith('.') and
            not re.match(r'^\d+', linha) and
            (linha.isupper() or linha.istitle())
        )
    
    def _eh_titulo(self, linha: str) -> bool:
        """Verifica se a linha é um título"""
        linha = linha.strip()
        
        # Padrões de títulos
        padroes_titulo = [
            r'^\d+\.\s+[A-Z]',  # "1. Título"
            r'^\d+\s+[A-Z]',    # "1 Título"
            r'^[A-Z][A-Z\s]{5,}$',  # Texto em maiúsculas
            r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*$'  # Título em formato de frase
        ]
        
        return (
            len(linha) > 3 and
            len(linha) < 80 and
            not linha.endswith('.') and
            any(re.match(padrao, linha) for padrao in padroes_titulo)
        )
    
    def _eh_subtitulo(self, linha: str) -> bool:
        """Verifica se a linha é um subtítulo"""
        linha = linha.strip()
        
        # Padrões de subtítulos
        padroes_subtitulo = [
            r'^\d+\.\d+\.?\s+[A-Z]',  # "1.1. Subtítulo"
            r'^\d+\.\d+\s+[A-Z]',     # "1.1 Subtítulo"
            r'^[a-z]\)\s+[A-Z]',       # "a) Subtítulo"
        ]
        
        return (
            len(linha) > 3 and
            len(linha) < 60 and
            not linha.endswith('.') and
            any(re.match(padrao, linha) for padrao in padroes_subtitulo)
        )
    
    def escrever_pdf(self, texto: str, arquivo_saida: str):
        """Escreve texto em formato PDF com formatação baseada na estrutura"""
        if not canvas:
            raise ImportError("reportlab não está instalado")
        
        estrutura = self._detectar_estrutura_documento(texto)
        
        doc = SimpleDocTemplate(arquivo_saida, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Estilos personalizados
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
        
        style_instituicao = ParagraphStyle(
            'Instituicao',
            parent=styles['Normal'],
            fontSize=12,
            alignment=TA_CENTER,
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        
        style_titulo_principal = ParagraphStyle(
            'TituloPrincipal',
            parent=styles['Title'],
            fontSize=16,
            alignment=TA_CENTER,
            spaceAfter=18,
            fontName='Helvetica-Bold'
        )
        
        style_secao_especial = ParagraphStyle(
            'SecaoEspecial',
            parent=styles['Heading1'],
            fontSize=14,
            alignment=TA_CENTER,
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        
        for item in estrutura:
            if item['tipo'] == 'instituicao':
                story.append(Paragraph(item['texto'], style_instituicao))
            elif item['tipo'] == 'titulo_principal':
                story.append(Paragraph(item['texto'], style_titulo_principal))
            elif item['tipo'] == 'secao_especial':
                story.append(Paragraph(item['texto'], style_secao_especial))
            elif item['tipo'] == 'titulo':
                story.append(Paragraph(item['texto'], styles['Heading1']))
            elif item['tipo'] == 'subtitulo':
                story.append(Paragraph(item['texto'], styles['Heading2']))
            elif item['tipo'] in ['lista_numerada', 'lista_marcador']:
                story.append(Paragraph(item['texto'], styles['Normal']))
            elif item['tipo'] == 'citacao':
                story.append(Paragraph(f'<i>{item["texto"]}</i>', styles['Normal']))
            elif item['tipo'] == 'referencia':
                story.append(Paragraph(item['texto'], styles['Normal']))
            else:  # paragrafo
                story.append(Paragraph(item['texto'], styles['Normal']))
            
            story.append(Spacer(1, 6))
        
        doc.build(story)
    
    def escrever_docx(self, texto: str, arquivo_saida: str):
        """Escreve texto em formato DOCX com formatação baseada na estrutura"""
        if not Document:
            raise ImportError("python-docx não está instalado")
        
        estrutura = self._detectar_estrutura_documento(texto)
        doc = Document()
        
        for item in estrutura:
            if item['tipo'] == 'instituicao':
                p = doc.add_heading(item['texto'], level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif item['tipo'] == 'titulo_principal':
                p = doc.add_heading(item['texto'], level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif item['tipo'] == 'secao_especial':
                p = doc.add_heading(item['texto'], level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif item['tipo'] == 'titulo':
                doc.add_heading(item['texto'], level=1)
            elif item['tipo'] == 'subtitulo':
                doc.add_heading(item['texto'], level=2)
            elif item['tipo'] in ['lista_numerada', 'lista_marcador']:
                doc.add_paragraph(item['texto'], style='List Bullet')
            elif item['tipo'] == 'citacao':
                p = doc.add_paragraph(item['texto'])
                p.style = 'Quote'
            elif item['tipo'] == 'referencia':
                p = doc.add_paragraph(item['texto'])
                p.style = 'Normal'
            else:  # paragrafo
                p = doc.add_paragraph(item['texto'])
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.save(arquivo_saida)
    
    def escrever_txt(self, texto: str, arquivo_saida: str):
        """Escreve texto em formato TXT preservando estrutura"""
        estrutura = self._detectar_estrutura_documento(texto)
        
        with open(arquivo_saida, 'w', encoding='utf-8') as arquivo:
            for item in estrutura:
                if item['tipo'] == 'instituicao':
                    arquivo.write(f"{item['texto'].center(80)}\n\n")
                elif item['tipo'] == 'titulo_principal':
                    arquivo.write(f"{item['texto'].center(80)}\n\n")
                elif item['tipo'] == 'secao_especial':
                    arquivo.write(f"\n{item['texto'].center(80)}\n\n")
                elif item['tipo'] == 'titulo':
                    arquivo.write(f"\n{item['texto']}\n{'=' * len(item['texto'])}\n\n")
                elif item['tipo'] == 'subtitulo':
                    arquivo.write(f"\n{item['texto']}\n{'-' * len(item['texto'])}\n\n")
                else:
                    arquivo.write(f"{item['texto']}\n")
    
    def escrever_html(self, texto: str, arquivo_saida: str):
        """Escreve texto em formato HTML com formatação baseada na estrutura"""
        estrutura = self._detectar_estrutura_documento(texto)
        
        html_content = """
<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Documento Convertido</title>
    <style>
        body {
            font-family: 'Times New Roman', serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f9f9f9;
        }
        .container {
            background-color: white;
            padding: 40px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }
        .instituicao {
            text-align: center;
            font-weight: bold;
            font-size: 14px;
            margin-bottom: 20px;
        }
        .titulo-principal {
            text-align: center;
            font-size: 18px;
            font-weight: bold;
            margin: 30px 0;
        }
        .secao-especial {
            text-align: center;
            font-size: 16px;
            font-weight: bold;
            margin: 25px 0;
        }
        h1 {
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 5px;
        }
        h2 {
            color: #34495e;
            margin-top: 25px;
        }
        p {
            text-align: justify;
            margin-bottom: 15px;
        }
        .citacao {
            font-style: italic;
            margin: 20px;
            padding: 10px;
            border-left: 3px solid #3498db;
            background-color: #f8f9fa;
        }
        ul, ol {
            margin-bottom: 15px;
        }
    </style>
</head>
<body>
    <div class="container">
"""
        
        for item in estrutura:
            if item['tipo'] == 'instituicao':
                html_content += f'        <div class="instituicao">{item["texto"]}</div>\n'
            elif item['tipo'] == 'titulo_principal':
                html_content += f'        <div class="titulo-principal">{item["texto"]}</div>\n'
            elif item['tipo'] == 'secao_especial':
                html_content += f'        <div class="secao-especial">{item["texto"]}</div>\n'
            elif item['tipo'] == 'titulo':
                html_content += f'        <h1>{item["texto"]}</h1>\n'
            elif item['tipo'] == 'subtitulo':
                html_content += f'        <h2>{item["texto"]}</h2>\n'
            elif item['tipo'] == 'lista_numerada':
                html_content += f'        <ol><li>{item["texto"][item["texto"].find(" ")+1:]}</li></ol>\n'
            elif item['tipo'] == 'lista_marcador':
                html_content += f'        <ul><li>{item["texto"][2:]}</li></ul>\n'
            elif item['tipo'] == 'citacao':
                html_content += f'        <div class="citacao">{item["texto"]}</div>\n'
            else:  # paragrafo
                html_content += f'        <p>{item["texto"]}</p>\n'
        
        html_content += """
    </div>
</body>
</html>
"""
        
        with open(arquivo_saida, 'w', encoding='utf-8') as arquivo:
            arquivo.write(html_content)
    
    def escrever_md(self, texto: str, arquivo_saida: str):
        """Escreve texto em formato Markdown com formatação baseada na estrutura"""
        estrutura = self._detectar_estrutura_documento(texto)
        
        with open(arquivo_saida, 'w', encoding='utf-8') as arquivo:
            for item in estrutura:
                if item['tipo'] == 'instituicao':
                    arquivo.write(f"<div align='center'>**{item['texto']}**</div>\n\n")
                elif item['tipo'] == 'titulo_principal':
                    arquivo.write(f"<div align='center'># {item['texto']}</div>\n\n")
                elif item['tipo'] == 'secao_especial':
                    arquivo.write(f"\n<div align='center'>## {item['texto']}</div>\n\n")
                elif item['tipo'] == 'titulo':
                    arquivo.write(f"\n# {item['texto']}\n\n")
                elif item['tipo'] == 'subtitulo':
                    arquivo.write(f"\n## {item['texto']}\n\n")
                elif item['tipo'] == 'lista_numerada':
                    arquivo.write(f"1. {item['texto'][item['texto'].find(' ')+1:]}\n")
                elif item['tipo'] == 'lista_marcador':
                    arquivo.write(f"- {item['texto'][2:]}\n")
                elif item['tipo'] == 'citacao':
                    arquivo.write(f"> {item['texto']}\n\n")
                else:  # paragrafo
                    arquivo.write(f"{item['texto']}\n\n")
    
    def converter(self, arquivo_origem: str, arquivo_destino: str, formato_destino: str = None) -> bool:
        """Converte um arquivo de um formato para outro"""
        try:
            # Detecta formato de origem
            formato_origem = self.detectar_formato(arquivo_origem)
            
            # Detecta formato de destino se não especificado
            if not formato_destino:
                formato_destino = self.detectar_formato(arquivo_destino)
            
            # Lê o arquivo de origem
            if formato_origem == 'pdf':
                texto = self.ler_pdf(arquivo_origem)
            elif formato_origem == 'docx':
                texto = self.ler_docx(arquivo_origem)
            elif formato_origem == 'txt':
                texto = self.ler_txt(arquivo_origem)
            elif formato_origem == 'html':
                texto = self.ler_html(arquivo_origem)
            elif formato_origem == 'md':
                texto = self.ler_md(arquivo_origem)
            else:
                raise ValueError(f"Formato de origem não suportado: {formato_origem}")
            
            # Escreve no formato de destino
            if formato_destino == 'pdf':
                self.escrever_pdf(texto, arquivo_destino)
            elif formato_destino == 'docx':
                self.escrever_docx(texto, arquivo_destino)
            elif formato_destino == 'txt':
                self.escrever_txt(texto, arquivo_destino)
            elif formato_destino == 'html':
                self.escrever_html(texto, arquivo_destino)
            elif formato_destino == 'md':
                self.escrever_md(texto, arquivo_destino)
            else:
                raise ValueError(f"Formato de destino não suportado: {formato_destino}")
            
            return True
            
        except Exception as e:
            print(f"Erro na conversão: {str(e)}")
            return False

# Configuração da aplicação Flask
app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max
app.config['UPLOAD_FOLDER'] = 'uploads'
app.secret_key = 'sua_chave_secreta_aqui'

# Configuração de CORS para permitir acesso do frontend
CORS(app, resources={r"/converter": {"origins": "http://localhost:3000"}})

# Cria diretório de upload se não existir
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Instância global do conversor
conversor = ConversorUniversalMelhorado()

def allowed_file(file_storage):
    """Verifica se a extensão e o tipo MIME do arquivo são permitidos."""
    if not file_storage or not file_storage.filename:
        return False

    filename = file_storage.filename
    allowed_ext = '.' in filename and \
                  filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

    if not allowed_ext:
        return False

    # Verifica o tipo MIME real do arquivo (se magic estiver disponível)
    if magic:
        file_storage.seek(0)
        mime_type = magic.from_buffer(file_storage.read(2048), mime=True)
        file_storage.seek(0) # Reseta o ponteiro novamente para o uso posterior
    else:
        # Se magic não estiver disponível, usa apenas verificação de extensão
        return True

    expected_mimes = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'doc': 'application/msword',
        'txt': 'text/plain',
        'html': 'text/html',
        'htm': 'text/html',
        'md': 'text/markdown',
        'markdown': 'text/markdown'
    }

    if magic:
        file_ext = filename.rsplit('.', 1)[1].lower()
        expected_mime = expected_mimes.get(file_ext)
        return mime_type == expected_mime
    else:
        # Se magic não estiver disponível, retorna True (já verificou extensão)
        return True

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/converter', methods=['POST'])
def converter_arquivo():
    try:
        if 'arquivo' not in request.files:
            return jsonify({'erro': 'Nenhum arquivo enviado'}), 400
        
        arquivo = request.files['arquivo']
        formato_destino = request.form.get('formato_destino')
        
        if arquivo.filename == '':
            return jsonify({'erro': 'Nenhum arquivo selecionado'}), 400

        if not formato_destino:
            return jsonify({'erro': 'Formato de destino não especificado'}), 400

        if not allowed_file(arquivo):
            return jsonify({'erro': 'Tipo de arquivo não permitido ou corrompido'}), 400
        
        # Salva arquivo temporário
        nome_arquivo = secure_filename(arquivo.filename)
        caminho_origem = os.path.join(app.config['UPLOAD_FOLDER'], nome_arquivo)
        arquivo.save(caminho_origem)
        
        # Define nome do arquivo de destino
        nome_base = os.path.splitext(nome_arquivo)[0]
        extensao_destino = conversor.formatos_suportados[formato_destino][0]
        nome_destino = f"{nome_base}_convertido{extensao_destino}"
        caminho_destino = os.path.join(app.config['UPLOAD_FOLDER'], nome_destino)
        
        # Realiza a conversão
        sucesso = conversor.converter(caminho_origem, caminho_destino, formato_destino)
        
        if sucesso:
            return send_file(caminho_destino, as_attachment=True, download_name=nome_destino)
        else:
            return jsonify({'erro': 'Falha na conversão'}), 500
            
    except Exception as e:
        return jsonify({'erro': f'Erro interno: {str(e)}'}), 500
    finally:
        # Limpa arquivos temporários
        try:
            if 'caminho_origem' in locals():
                os.remove(caminho_origem)
        except:
            pass

@app.route('/formatos')
def listar_formatos():
    return jsonify({
        'formatos_suportados': list(conversor.formatos_suportados.keys()),
        'detalhes': conversor.formatos_suportados
    })

# Template HTML melhorado
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Conversor Universal de Documentos - Melhorado</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            display: flex;
            align-items: center;
            justify-content: center;
        }
        
        .container {
            background: white;
            padding: 2rem;
            border-radius: 15px;
            box-shadow: 0 20px 40px rgba(0,0,0,0.1);
            max-width: 500px;
            width: 90%;
        }
        
        h1 {
            text-align: center;
            color: #333;
            margin-bottom: 1.5rem;
            font-size: 1.8rem;
        }
        
        .subtitle {
            text-align: center;
            color: #666;
            margin-bottom: 2rem;
            font-size: 0.9rem;
        }
        
        .upload-area {
            border: 2px dashed #ddd;
            border-radius: 10px;
            padding: 2rem;
            text-align: center;
            margin-bottom: 1.5rem;
            transition: all 0.3s ease;
            cursor: pointer;
        }
        
        .upload-area:hover {
            border-color: #667eea;
            background-color: #f8f9ff;
        }
        
        .upload-area.dragover {
            border-color: #667eea;
            background-color: #f0f4ff;
        }
        
        input[type="file"] {
            display: none;
        }
        
        .file-info {
            margin: 1rem 0;
            padding: 0.5rem;
            background-color: #f8f9fa;
            border-radius: 5px;
            font-size: 0.9rem;
        }
        
        select {
            width: 100%;
            padding: 0.75rem;
            border: 1px solid #ddd;
            border-radius: 5px;
            font-size: 1rem;
            margin-bottom: 1.5rem;
        }
        
        button {
            width: 100%;
            padding: 0.75rem;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            border-radius: 5px;
            font-size: 1rem;
            cursor: pointer;
            transition: transform 0.2s ease;
        }
        
        button:hover {
            transform: translateY(-2px);
        }
        
        button:disabled {
            background: #ccc;
            cursor: not-allowed;
            transform: none;
        }
        
        .progress {
            width: 100%;
            height: 4px;
            background-color: #f0f0f0;
            border-radius: 2px;
            margin: 1rem 0;
            overflow: hidden;
            display: none;
        }
        
        .progress-bar {
            height: 100%;
            background: linear-gradient(90deg, #667eea, #764ba2);
            width: 0%;
            transition: width 0.3s ease;
        }
        
        .message {
            padding: 0.75rem;
            border-radius: 5px;
            margin: 1rem 0;
            display: none;
        }
        
        .message.success {
            background-color: #d4edda;
            color: #155724;
            border: 1px solid #c3e6cb;
        }
        
        .message.error {
            background-color: #f8d7da;
            color: #721c24;
            border: 1px solid #f5c6cb;
        }
        
        .formats-info {
            background-color: #e7f3ff;
            padding: 1rem;
            border-radius: 5px;
            margin-bottom: 1.5rem;
            font-size: 0.85rem;
        }
        
        .formats-info h3 {
            margin-bottom: 0.5rem;
            color: #0066cc;
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>🚀 Conversor Universal Melhorado</h1>
        <p class="subtitle">Converta documentos preservando a estrutura original</p>
        
        <div class="formats-info">
            <h3>📋 Formatos Suportados:</h3>
            <p><strong>Entrada:</strong> PDF, DOCX, TXT, HTML, MD</p>
            <p><strong>Saída:</strong> PDF, DOCX, TXT, HTML, MD</p>
            <p><strong>✨ Novo:</strong> Preservação inteligente de estrutura acadêmica!</p>
        </div>
        
        <form id="converterForm" enctype="multipart/form-data">
            <div class="upload-area" onclick="document.getElementById('arquivo').click()">
                <p>📁 Clique aqui ou arraste um arquivo</p>
                <p style="font-size: 0.8rem; color: #666; margin-top: 0.5rem;">Máximo: 16MB</p>
                <input type="file" id="arquivo" name="arquivo" accept=".pdf,.docx,.doc,.txt,.html,.htm,.md,.markdown">
            </div>
            
            <div class="file-info" id="fileInfo" style="display: none;"></div>
            
            <select id="formatoDestino" name="formato_destino" required>
                <option value="">Selecione o formato de destino</option>
                <option value="pdf">PDF</option>
                <option value="docx">DOCX (Word)</option>
                <option value="txt">TXT (Texto)</option>
                <option value="html">HTML</option>
                <option value="md">Markdown</option>
            </select>
            
            <button type="submit" id="btnConverter" disabled>🔄 Converter Documento</button>
            
            <div class="progress" id="progress">
                <div class="progress-bar" id="progressBar"></div>
            </div>
            
            <div class="message" id="message"></div>
        </form>
    </div>
    
    <script>
        const form = document.getElementById('converterForm');
        const fileInput = document.getElementById('arquivo');
        const fileInfo = document.getElementById('fileInfo');
        const btnConverter = document.getElementById('btnConverter');
        const progress = document.getElementById('progress');
        const progressBar = document.getElementById('progressBar');
        const message = document.getElementById('message');
        const uploadArea = document.querySelector('.upload-area');
        
        // Drag and drop
        uploadArea.addEventListener('dragover', (e) => {
            e.preventDefault();
            uploadArea.classList.add('dragover');
        });
        
        uploadArea.addEventListener('dragleave', () => {
            uploadArea.classList.remove('dragover');
        });
        
        uploadArea.addEventListener('drop', (e) => {
            e.preventDefault();
            uploadArea.classList.remove('dragover');
            const files = e.dataTransfer.files;
            if (files.length > 0) {
                fileInput.files = files;
                updateFileInfo();
            }
        });
        
        fileInput.addEventListener('change', updateFileInfo);
        
        function updateFileInfo() {
            const file = fileInput.files[0];
            if (file) {
                const size = (file.size / 1024 / 1024).toFixed(2);
                fileInfo.innerHTML = `
                    <strong>📄 Arquivo selecionado:</strong><br>
                    Nome: ${file.name}<br>
                    Tamanho: ${size} MB
                `;
                fileInfo.style.display = 'block';
                btnConverter.disabled = false;
            } else {
                fileInfo.style.display = 'none';
                btnConverter.disabled = true;
            }
        }
        
        form.addEventListener('submit', async (e) => {
            e.preventDefault();
            
            const formData = new FormData(form);
            
            // Mostra progresso
            progress.style.display = 'block';
            btnConverter.disabled = true;
            btnConverter.textContent = '🔄 Convertendo...';
            
            // Simula progresso
            let progressValue = 0;
            const progressInterval = setInterval(() => {
                progressValue += Math.random() * 30;
                if (progressValue > 90) progressValue = 90;
                progressBar.style.width = progressValue + '%';
            }, 200);
            
            try {
                const response = await fetch('/converter', {
                    method: 'POST',
                    body: formData
                });
                
                clearInterval(progressInterval);
                progressBar.style.width = '100%';
                
                if (response.ok) {
                    const blob = await response.blob();
                    const url = window.URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = response.headers.get('Content-Disposition')?.split('filename=')[1] || 'arquivo_convertido';
                    document.body.appendChild(a);
                    a.click();
                    window.URL.revokeObjectURL(url);
                    document.body.removeChild(a);
                    
                    showMessage('✅ Conversão realizada com sucesso! Download iniciado.', 'success');
                } else {
                    const error = await response.json();
                    showMessage(`❌ Erro: ${error.erro}`, 'error');
                }
            } catch (error) {
                clearInterval(progressInterval);
                showMessage(`❌ Erro de conexão: ${error.message}`, 'error');
            } finally {
                progress.style.display = 'none';
                progressBar.style.width = '0%';
                btnConverter.disabled = false;
                btnConverter.textContent = '🔄 Converter Documento';
            }
        });
        
        function showMessage(text, type) {
            message.textContent = text;
            message.className = `message ${type}`;
            message.style.display = 'block';
            
            setTimeout(() => {
                message.style.display = 'none';
            }, 5000);
        }
    </script>
</body>
</html>
"""

if __name__ == '__main__':
    print("🚀 Iniciando Conversor Universal de Documentos Melhorado...")
    print("📋 Formatos suportados:")
    for formato, extensoes in conversor.formatos_suportados.items():
        print(f"   {formato.upper()}: {', '.join(extensoes)}")
    
    print("\n🔧 Dependências opcionais:")
    print(f"   📄 PDF: {'✅' if pdfplumber else '❌'} pdfplumber")
    print(f"   📝 DOCX: {'✅' if Document else '❌'} python-docx")
    print(f"   🎨 PDF Output: {'✅' if canvas else '❌'} reportlab")
    print(f"   🌐 HTML: {'✅' if BeautifulSoup else '❌'} beautifulsoup4")
    print(f"   📖 Markdown: {'✅' if markdown else '❌'} markdown")
    
    print("\n🌟 Melhorias implementadas:")
    print("   ✨ Detecção inteligente de estrutura acadêmica")
    print("   📚 Preservação de títulos, subtítulos e seções")
    print("   🎯 Formatação específica por tipo de documento")
    print("   🔧 Correção automática de espaçamento")
    
    print("\n🌐 Servidor iniciando...")
    app.run(debug=True, host='0.0.0.0', port=5000)

def create_app(config_name: str = None) -> Flask:
    """
    Application factory pattern for creating Flask app.
    
    Args:
        config_name: Configuration environment name
        
    Returns:
        Configured Flask application
    """
    # Create Flask app
    app = Flask(__name__)
    
    # Load configuration
    config_class = get_config(config_name)
    app.config.from_object(config_class)
    config_class.init_app(app)
    
    # Setup logging
    setup_logging(
        log_level=app.config.get('LOG_LEVEL', 'INFO'),
        log_dir=Path(app.config.get('LOG_DIR', 'logs')),
        app_name='conversor',
        enable_json=app.config.get('ENABLE_JSON_LOGGING', False)
    )
    
    # Setup CORS
    CORS(app, resources={
        r"/api/*": {
            "origins": app.config.get('ALLOWED_ORIGINS', ["http://localhost:3000"]),
            "methods": ["GET", "POST", "OPTIONS"],
            "allow_headers": ["Content-Type", "Authorization"]
        }
    })
    
    # Register blueprints
    app.register_blueprint(api_bp)
    
    # Initialize rate limiter
    app.rate_limiter = IPRateLimiter()
    
    # Health check endpoint
    @app.route('/health')
    def health_check():
        return {'status': 'healthy', 'version': '2.0.0'}
    
    # Error handlers
    @app.errorhandler(404)
    def not_found(error):
        return {'error': 'Not found'}, 404
    
    @app.errorhandler(500)
    def internal_error(error):
        return {'error': 'Internal server error'}, 500
    
    return app


# Create app instance
app = create_app()

if __name__ == '__main__':
    # Development server
    app.run(
        host=app.config.get('HOST', '127.0.0.1'),
        port=app.config.get('PORT', 5000),
        debug=app.config.get('DEBUG', False)
    )